from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


MAX_UPLOAD_SIZE_BYTES = 10 * 1024 * 1024


def extract_text_from_cv_file(filename: str, data: bytes) -> str:
    if not data:
        raise ValueError("Uploaded CV file is empty")

    if len(data) > MAX_UPLOAD_SIZE_BYTES:
        raise ValueError("Uploaded CV file is too large. Max size is 10 MB")

    suffix = Path(filename or "").suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf_text(data)
    elif suffix == ".docx":
        text = _extract_docx_text(data)
    elif suffix == ".txt":
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError("Unsupported CV file type. Use PDF, DOCX, TXT, or JSON")

    normalized = "\n".join(line.strip() for line in text.splitlines() if line.strip())

    if not normalized:
        raise ValueError("Could not extract text from the uploaded CV file")

    return normalized


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    document = Document(BytesIO(data))
    parts: list[str] = []

    parts.extend(paragraph.text for paragraph in document.paragraphs)

    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts)
